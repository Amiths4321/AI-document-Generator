# doc_builder.py
import os
from pathlib  import Path
from datetime import datetime

OUTPUT_DIR = Path("generated_docs")
OUTPUT_DIR.mkdir(exist_ok=True)


def build_docx(content: str, doc_type: str, title: str = "") -> str:
    """Build a formatted Word document from generated content."""
    from docx                   import Document
    from docx.shared            import Pt, Inches, RGBColor
    from docx.enum.text         import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns           import qn
    from docx.oxml              import OxmlElement

    doc   = Document()
    style = doc.styles["Normal"]
    font  = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title_text = title or doc_type
    heading    = doc.add_heading(title_text, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date and metadata
    meta = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.color.rgb = RGBColor(0x88, 0x87, 0x80)

    doc.add_paragraph()   # spacer

    # Process content line by line
    lines = content.split("\n")
    i     = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            doc.add_paragraph()
            i += 1
            continue

        # Detect headings (numbered sections or ALL CAPS)
        if (line[0].isdigit() and ". " in line[:5]) or line.isupper():
            h = doc.add_heading(line, level=2)
            h.runs[0].font.color.rgb = RGBColor(0x1D, 0x9E, 0x75)

        # Bullet points
        elif line.startswith("- ") or line.startswith("• "):
            p = doc.add_paragraph(
                line.lstrip("- •"),
                style="List Bullet"
            )

        # Bold terms (surrounded by **)
        elif "**" in line:
            p = doc.add_paragraph()
            parts = line.split("**")
            for j, part in enumerate(parts):
                run = p.add_run(part)
                if j % 2 == 1:
                    run.bold = True

        # Signature lines
        elif "_____" in line or "Signature" in line:
            p = doc.add_paragraph(line)
            p.runs[0].font.color.rgb = RGBColor(0x5F, 0x5E, 0x5A)

        # Normal paragraph
        else:
            doc.add_paragraph(line)

        i += 1

    # Save
    ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe     = "".join(c if c.isalnum() else "_" for c in doc_type[:30])
    filepath = OUTPUT_DIR / f"{safe}_{ts}.docx"
    doc.save(str(filepath))
    return str(filepath)


def build_txt(content: str, doc_type: str) -> str:
    """Save as plain text file."""
    ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe     = "".join(c if c.isalnum() else "_" for c in doc_type[:30])
    filepath = OUTPUT_DIR / f"{safe}_{ts}.txt"
    filepath.write_text(content, encoding="utf-8")
    return str(filepath)


def build_markdown(content: str, doc_type: str) -> str:
    """Save as markdown file."""
    ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe     = "".join(c if c.isalnum() else "_" for c in doc_type[:30])
    filepath = OUTPUT_DIR / f"{safe}_{ts}.md"
    header   = f"# {doc_type}\n*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n"
    filepath.write_text(header + content, encoding="utf-8")
    return str(filepath)